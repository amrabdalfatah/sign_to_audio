from docx import Document
from fpdf import FPDF

# Create Word document
doc = Document()
doc.add_heading('Linear Regression Homework Solution', 0)

doc.add_paragraph("""Using linear regression matrix:

[ [n        Σc ]
  [Σc    Σc²] ] × [ b ]
                [ a ] = [ Σk ]
                          [ Σck ]

Find summations:

| N | c   | k   | c²    | c * k  |
|---|-----|-----|--------|--------|
| 1 | 0.8 | 1   | 0.64   | 0.8    |
| 2 | 2   | 5.6 | 4      | 11.2   |
| 3 | 3.6 | 7.2 | 12.96  | 25.92  |
| 4 | 5   | 10  | 25     | 50     |
|   | 11.4| 23.8| 42.6   | 87.92  |

Solve linear regression matrix using calculator:

[ [4     11.4 ]
  [11.4  42.6] ] × [ b ]
                  [ a ] = [ 23.8 ]
                           [ 87.92 ]

b = 0.2866
a = 1.9871

The equation of straight line will be:
k = 1.9871c + 0.2866

To calculate St, Sᵧ/ₓ, r²:

k̄ = Σk / n = 23.8 / 4 = 5.95

| N | (k - k̄)² | (k - ac - b)² |
|---|----------|----------------|
| 1 | 24.5025  | 0.768          |
| 2 | 0.1225   | 1.7931         |
| 3 | 1.5625   | 0.0578         |
| 4 | 16.4025  | 0.0494         |
|   | ∑ = 42.59| ∑ = 2.668      |

Then:

s_y = sqrt(St / (n - 1)) = sqrt(42.59 / 3) = 3.767

s_y/x = sqrt(Sr / (n - 2)) = sqrt(2.668 / 2) = 1.155

r² = (St - Sr) / St = (42.59 - 2.668) / 42.59 = 0.937
""")

word_path = "/mnt/data/Linear_Regression_Solution.docx"
doc.save(word_path)

# Create PDF
pdf = FPDF()
pdf.add_page()
pdf.set_auto_page_break(auto=True, margin=15)
pdf.set_font("Arial", size=12)
pdf.multi_cell(0, 10, doc.paragraphs[1].text)

pdf_path = "/mnt/data/Linear_Regression_Solution.pdf"
pdf.output(pdf_path)

word_path, pdf_path
